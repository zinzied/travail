"""
Document text extraction utilities for PDF, Word, and Excel files.
"""

import logging
from pathlib import Path
from typing import Optional, Dict, Any, List
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
from ..utils.exceptions import PDFExtractionError

logger = logging.getLogger(__name__)


class DocumentExtractionError(Exception):
    """Exception raised when document extraction fails."""
    pass


class PDFExtractor:
    """Extract text and metadata from PDF files using multiple methods."""
    
    def __init__(self):
        self.extraction_methods = [
            self._extract_with_pdfplumber,
            self._extract_with_pymupdf,
            self._extract_with_pypdf2
        ]
    
    def extract_text(self, pdf_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF using the best available method.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        pdf_path = Path(pdf_path)
        
        if not pdf_path.exists():
            raise PDFExtractionError(f"PDF file not found: {pdf_path}")
        
        if not pdf_path.suffix.lower() == '.pdf':
            raise PDFExtractionError(f"File is not a PDF: {pdf_path}")
        
        # Try extraction methods in order of preference
        for method in self.extraction_methods:
            try:
                result = method(pdf_path)
                if result and result.get('text', '').strip():
                    logger.info(f"Successfully extracted text using {method.__name__}")
                    return result
            except Exception as e:
                logger.warning(f"Failed to extract with {method.__name__}: {e}")
                continue
        
        raise PDFExtractionError(f"Failed to extract text from PDF: {pdf_path}")
    
    def _extract_with_pdfplumber(self, pdf_path: Path) -> Dict[str, Any]:
        """Extract text using pdfplumber (best for structured documents)."""
        text_content = []
        metadata = {}
        
        with pdfplumber.open(pdf_path) as pdf:
            metadata = {
                'pages': len(pdf.pages),
                'method': 'pdfplumber',
                'file_size': pdf_path.stat().st_size
            }
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        return {
            'text': '\n'.join(text_content),
            'metadata': metadata,
            'confidence': 0.9 if text_content else 0.0
        }
    
    def _extract_with_pymupdf(self, pdf_path: Path) -> Dict[str, Any]:
        """Extract text using PyMuPDF (good for complex layouts)."""
        text_content = []
        
        doc = fitz.open(pdf_path)
        metadata = {
            'pages': doc.page_count,
            'method': 'pymupdf',
            'file_size': pdf_path.stat().st_size,
            'pdf_metadata': doc.metadata
        }
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text:
                text_content.append(page_text)
        
        doc.close()
        
        return {
            'text': '\n'.join(text_content),
            'metadata': metadata,
            'confidence': 0.8 if text_content else 0.0
        }
    
    def _extract_with_pypdf2(self, pdf_path: Path) -> Dict[str, Any]:
        """Extract text using PyPDF2 (fallback method)."""
        text_content = []
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            metadata = {
                'pages': len(pdf_reader.pages),
                'method': 'pypdf2',
                'file_size': pdf_path.stat().st_size
            }
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        return {
            'text': '\n'.join(text_content),
            'metadata': metadata,
            'confidence': 0.7 if text_content else 0.0
        }


class WordExtractor:
    """Extract text and metadata from Word documents (.docx)."""

    def extract_text(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract text from Word document.

        Args:
            docx_path: Path to the Word document

        Returns:
            Dictionary containing extracted text and metadata
        """
        docx_path = Path(docx_path)

        if not docx_path.exists():
            raise DocumentExtractionError(f"Word document not found: {docx_path}")

        if not docx_path.suffix.lower() in ['.docx', '.doc']:
            raise DocumentExtractionError(f"File is not a Word document: {docx_path}")

        try:
            doc = Document(docx_path)

            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))

            # Combine all text
            all_text = []
            if paragraphs:
                all_text.extend(paragraphs)
            if table_text:
                all_text.append("\n--- Tables ---")
                all_text.extend(table_text)

            text_content = '\n'.join(all_text)

            # Metadata
            metadata = {
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
                'method': 'python-docx',
                'file_size': docx_path.stat().st_size
            }

            return {
                'text': text_content,
                'metadata': metadata,
                'confidence': 0.9 if text_content.strip() else 0.0
            }

        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {e}")
            raise DocumentExtractionError(f"Failed to extract text from Word document: {docx_path}")


class ExcelExtractor:
    """Extract text and data from Excel files (.xlsx, .xls)."""

    def extract_text(self, excel_path: str) -> Dict[str, Any]:
        """
        Extract text from Excel file.

        Args:
            excel_path: Path to the Excel file

        Returns:
            Dictionary containing extracted text and metadata
        """
        excel_path = Path(excel_path)

        if not excel_path.exists():
            raise DocumentExtractionError(f"Excel file not found: {excel_path}")

        if not excel_path.suffix.lower() in ['.xlsx', '.xls']:
            raise DocumentExtractionError(f"File is not an Excel document: {excel_path}")

        try:
            # Read all sheets
            excel_file = pd.ExcelFile(excel_path)
            sheet_names = excel_file.sheet_names

            all_text = []
            total_rows = 0
            total_cols = 0

            for sheet_name in sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)

                if not df.empty:
                    all_text.append(f"\n--- Sheet: {sheet_name} ---")

                    # Add column headers
                    headers = ' | '.join(str(col) for col in df.columns)
                    all_text.append(f"Headers: {headers}")

                    # Add data rows (limit to first 100 rows for performance)
                    for idx, row in df.head(100).iterrows():
                        row_text = ' | '.join(str(val) if pd.notna(val) else '' for val in row.values)
                        if row_text.strip():
                            all_text.append(row_text)

                    total_rows += len(df)
                    total_cols = max(total_cols, len(df.columns))

                    if len(df) > 100:
                        all_text.append(f"... ({len(df) - 100} more rows)")

            text_content = '\n'.join(all_text)

            # Metadata
            metadata = {
                'sheets': len(sheet_names),
                'total_rows': total_rows,
                'total_columns': total_cols,
                'sheet_names': sheet_names,
                'method': 'pandas',
                'file_size': excel_path.stat().st_size
            }

            return {
                'text': text_content,
                'metadata': metadata,
                'confidence': 0.8 if text_content.strip() else 0.0
            }

        except Exception as e:
            logger.error(f"Failed to extract text from Excel file: {e}")
            raise DocumentExtractionError(f"Failed to extract text from Excel file: {excel_path}")


class UniversalDocumentExtractor:
    """Universal document extractor that can handle PDF, Word, and Excel files."""

    def __init__(self):
        self.pdf_extractor = PDFExtractor()
        self.word_extractor = WordExtractor()
        self.excel_extractor = ExcelExtractor()

        # Supported file extensions
        self.supported_extensions = {
            '.pdf': self.pdf_extractor,
            '.docx': self.word_extractor,
            '.doc': self.word_extractor,
            '.xlsx': self.excel_extractor,
            '.xls': self.excel_extractor
        }

    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from any supported document type.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise DocumentExtractionError(f"File not found: {file_path}")

        file_extension = file_path.suffix.lower()

        if file_extension not in self.supported_extensions:
            raise DocumentExtractionError(
                f"Unsupported file type: {file_extension}. "
                f"Supported types: {', '.join(self.supported_extensions.keys())}"
            )

        try:
            extractor = self.supported_extensions[file_extension]
            result = extractor.extract_text(str(file_path))

            # Add file type to metadata
            result['metadata']['file_type'] = file_extension
            result['metadata']['file_name'] = file_path.name

            logger.info(f"Successfully extracted text from {file_extension} file: {file_path.name}")
            return result

        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise DocumentExtractionError(f"Failed to extract text from {file_path}: {e}")

    def is_supported(self, file_path: str) -> bool:
        """Check if file type is supported."""
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.supported_extensions

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.supported_extensions.keys())

    def get_file_type_description(self, file_path: str) -> str:
        """Get human-readable description of file type."""
        file_extension = Path(file_path).suffix.lower()

        descriptions = {
            '.pdf': 'PDF Document',
            '.docx': 'Word Document',
            '.doc': 'Word Document (Legacy)',
            '.xlsx': 'Excel Spreadsheet',
            '.xls': 'Excel Spreadsheet (Legacy)'
        }

        return descriptions.get(file_extension, 'Unknown File Type')
    
    def validate_pdf(self, pdf_path: str) -> bool:
        """
        Validate if the file is a readable PDF.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            True if valid PDF, False otherwise
        """
        try:
            result = self.extract_text(pdf_path)
            return bool(result.get('text', '').strip())
        except Exception:
            return False
