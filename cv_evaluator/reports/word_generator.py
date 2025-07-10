"""
Word document report generator using python-docx.
"""

import logging
from pathlib import Path
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from ..utils.exceptions import ReportGenerationError

logger = logging.getLogger(__name__)


class WordReportGenerator:
    """Generate Word document evaluation reports using python-docx."""
    
    def __init__(self, template_dir: str = "templates"):
        self.template_dir = Path(template_dir)
    
    def generate(self, report_data: Dict[str, Any], output_path: str, 
                template: str = "default", **kwargs) -> str:
        """
        Generate Word document report.
        
        Args:
            report_data: Report data dictionary
            output_path: Output file path
            template: Template name (not used in this implementation)
            **kwargs: Additional options
            
        Returns:
            Path to generated Word document
        """
        try:
            output_path = Path(output_path)
            if output_path.suffix.lower() not in ['.docx', '.doc']:
                output_path = output_path.with_suffix('.docx')
            
            # Create document
            doc = Document()
            
            # Set document styles
            self._setup_document_styles(doc)
            
            # Build content
            self._add_header(doc, report_data)
            self._add_summary(doc, report_data)
            self._add_section_scores(doc, report_data)
            self._add_detailed_analysis(doc, report_data)
            self._add_skills_analysis(doc, report_data)
            self._add_experience_analysis(doc, report_data)
            self._add_education_analysis(doc, report_data)
            self._add_recommendations(doc, report_data)
            
            # Save document
            doc.save(str(output_path))
            
            logger.info(f"Word report generated: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Word generation failed: {e}")
            raise ReportGenerationError(f"Failed to generate Word report: {e}")
    
    def _setup_document_styles(self, doc: Document):
        """Setup document styles."""
        # Title style
        title_style = doc.styles['Title']
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.color.rgb = RGBColor(44, 90, 160)  # Blue color
        
        # Heading styles
        heading1_style = doc.styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Arial'
        heading1_font.size = Pt(16)
        heading1_font.color.rgb = RGBColor(44, 90, 160)
        
        # Normal style
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
    
    def _add_header(self, doc: Document, report_data: Dict[str, Any]):
        """Add document header."""
        metadata = report_data['metadata']
        personal_info = report_data['personal_info']
        
        # Title
        title = doc.add_heading('CV Evaluation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Candidate name
        if personal_info.get('name'):
            name_para = doc.add_heading(personal_info['name'], level=2)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generation info
        gen_date = metadata['generated_at'].strftime('%Y-%m-%d %H:%M:%S')
        gen_para = doc.add_paragraph(f"Generated on: {gen_date}")
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add space
        doc.add_paragraph()
    
    def _add_summary(self, doc: Document, report_data: Dict[str, Any]):
        """Add executive summary."""
        summary = report_data['summary']
        
        doc.add_heading('Executive Summary', level=1)
        
        # Overall score
        score_para = doc.add_paragraph()
        score_run = score_para.add_run(f"Overall Score: {summary['overall_score']:.1f}/100")
        score_run.font.size = Pt(18)
        score_run.font.bold = True
        score_run.font.color.rgb = RGBColor(44, 90, 160)
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fit percentage
        fit_para = doc.add_paragraph()
        fit_run = fit_para.add_run(f"Job Fit: {summary['fit_percentage']:.1f}%")
        fit_run.font.size = Pt(18)
        fit_run.font.bold = True
        fit_run.font.color.rgb = RGBColor(44, 90, 160)
        fit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary table
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        
        # Data rows
        data_rows = [
            ('Strengths Identified', str(summary['strengths_count'])),
            ('Areas for Improvement', str(summary['weaknesses_count'])),
            ('Recommendations', str(summary['recommendations_count']))
        ]
        
        for i, (metric, value) in enumerate(data_rows, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        doc.add_paragraph()
    
    def _add_section_scores(self, doc: Document, report_data: Dict[str, Any]):
        """Add section scores breakdown."""
        doc.add_heading('Section Scores', level=1)
        
        # Create scores table
        table = doc.add_table(rows=len(report_data['section_scores']) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Section'
        header_cells[1].text = 'Score'
        header_cells[2].text = 'Percentage'
        header_cells[3].text = 'Feedback'
        
        # Data rows
        for i, score in enumerate(report_data['section_scores'], 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = score['section']
            row_cells[1].text = f"{score['score']:.1f}/{score['max_score']:.1f}"
            row_cells[2].text = f"{score['percentage']:.1f}%"
            row_cells[3].text = score['feedback'][:100] + "..." if len(score['feedback']) > 100 else score['feedback']
        
        doc.add_paragraph()
    
    def _add_detailed_analysis(self, doc: Document, report_data: Dict[str, Any]):
        """Add detailed analysis section."""
        analysis = report_data['analysis']
        
        doc.add_heading('Detailed Analysis', level=1)
        
        # Strengths
        if analysis['strengths']:
            doc.add_heading('Strengths:', level=2)
            for strength in analysis['strengths']:
                para = doc.add_paragraph(strength, style='List Bullet')
        
        # Weaknesses
        if analysis['weaknesses']:
            doc.add_heading('Areas for Improvement:', level=2)
            for weakness in analysis['weaknesses']:
                para = doc.add_paragraph(weakness, style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_skills_analysis(self, doc: Document, report_data: Dict[str, Any]):
        """Add skills analysis section."""
        skills = report_data['skills']
        
        if not skills:
            return
        
        doc.add_heading('Skills Analysis', level=1)
        
        # Create skills table
        table = doc.add_table(rows=min(len(skills), 15) + 1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Skill'
        header_cells[1].text = 'Category'
        header_cells[2].text = 'Level'
        header_cells[3].text = 'Experience'
        header_cells[4].text = 'Confidence'
        
        # Data rows (limit to top 15 skills)
        for i, skill in enumerate(skills[:15], 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = skill['name']
            row_cells[1].text = skill['category']
            row_cells[2].text = skill['level']
            row_cells[3].text = f"{skill['years_experience']} years" if skill['years_experience'] else "N/A"
            row_cells[4].text = f"{skill['confidence']:.1f}"
        
        doc.add_paragraph()
    
    def _add_experience_analysis(self, doc: Document, report_data: Dict[str, Any]):
        """Add work experience analysis."""
        experience = report_data['experience']
        
        if not experience:
            return
        
        doc.add_heading('Work Experience Analysis', level=1)
        
        for exp in experience:
            # Company and position
            exp_heading = f"{exp['position']} at {exp['company']}"
            doc.add_heading(exp_heading, level=3)
            
            # Duration
            if exp['start_date'] and exp['end_date']:
                duration = f"{exp['start_date']} - {exp['end_date']}"
                if exp['duration_months']:
                    duration += f" ({exp['duration_months']} months)"
                doc.add_paragraph(duration)
            
            # Technologies
            if exp['technologies']:
                tech_para = doc.add_paragraph()
                tech_para.add_run("Technologies: ").bold = True
                tech_para.add_run(', '.join(exp['technologies'][:10]))
            
            # Key achievements
            if exp['achievements']:
                doc.add_paragraph("Key Achievements:").runs[0].bold = True
                for achievement in exp['achievements'][:3]:
                    doc.add_paragraph(achievement, style='List Bullet')
            
            doc.add_paragraph()
    
    def _add_education_analysis(self, doc: Document, report_data: Dict[str, Any]):
        """Add education analysis."""
        education = report_data['education']
        
        if not education:
            return
        
        doc.add_heading('Education Analysis', level=1)
        
        # Education table
        table = doc.add_table(rows=len(education) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Institution'
        header_cells[1].text = 'Degree'
        header_cells[2].text = 'Field of Study'
        header_cells[3].text = 'Year'
        
        # Data rows
        for i, edu in enumerate(education, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = edu['institution']
            row_cells[1].text = edu['degree']
            row_cells[2].text = edu['field_of_study'] or 'N/A'
            row_cells[3].text = str(edu['graduation_year']) if edu['graduation_year'] else 'N/A'
        
        doc.add_paragraph()
    
    def _add_recommendations(self, doc: Document, report_data: Dict[str, Any]):
        """Add recommendations section."""
        recommendations = report_data['analysis']['recommendations']
        
        if not recommendations:
            return
        
        doc.add_heading('Recommendations', level=1)
        
        for i, recommendation in enumerate(recommendations, 1):
            para = doc.add_paragraph()
            para.add_run(f"{i}. ").bold = True
            para.add_run(recommendation)
        
        doc.add_paragraph()
